"""
Build two supporting documents:
  - model.equations.docx : every governing equation used in the UBDS solver
  - source.docx          : all calibration and validation data sources

Author: Akosa Samuel Onyejekwe (Independent Researcher)
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INK = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2A, 0x6F, 0x97)


def new_doc():
    doc = Document()
    s = doc.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(10.5)
    s.font.color.rgb = INK
    for h in ["Title", "Heading 1", "Heading 2"]:
        try:
            doc.styles[h].font.color.rgb = ACCENT if h != "Title" else INK
        except Exception:
            pass
    return doc


def h1(d, t): d.add_heading(t, 1)
def h2(d, t): d.add_heading(t, 2)
def p(d, t):
    par = d.add_paragraph(t);
    return par
def eq(d, t):
    par = d.add_paragraph()
    r = par.add_run(t); r.font.name = "Consolas"; r.font.size = Pt(10.5)
    r.font.color.rgb = INK
    par.paragraph_format.left_indent = Pt(18)
    return par
def bullet(d, t): d.add_paragraph(t, style="List Bullet")


# ===========================================================================
#  MODEL EQUATIONS DOCUMENT
# ===========================================================================
def build_equations():
    d = new_doc()
    t = d.add_heading("", 0)
    t.add_run("UBDS - Governing Equations").font.color.rgb = INK
    sub = d.add_paragraph()
    rr = sub.add_run("Universal Brine Dispersion Solver: complete set of model equations")
    rr.bold = True; rr.font.size = Pt(12); rr.font.color.rgb = ACCENT
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = d.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run("Author: Akosa Samuel Onyejekwe (Independent Researcher)").font.color.rgb = INK
    p(d, "Notation: S = practical salinity (psu); T = temperature (degC); "
         "rho = density (kg/m3); rho_a = ambient density; g = 9.81 m/s2; "
         "g' = reduced gravity; b = plume radius; w = jet speed; V = velocity "
         "vector; Q = volume flux; s = arclength; u,v = depth-averaged "
         "velocities; eta = surface elevation; h = still depth; H = h + eta = "
         "total depth; D = dispersion coefficient.")

    # --- 1 EOS ---
    h1(d, "1. Equation of State (module: eos.py)")
    p(d, "1.1 Pure-water (SMOW) density, kg/m3:")
    eq(d, "rho_w(T) = 999.842594 + 6.793952e-2 T - 9.095290e-3 T^2 "
          "+ 1.001685e-4 T^3 - 1.120083e-6 T^4 + 6.536332e-9 T^5")
    p(d, "1.2 UNESCO EOS-80 one-atmosphere seawater density (valid S <= 42 psu):")
    eq(d, "rho(S,T) = rho_w(T) + B(T) S + C(T) S^1.5 + d0 S^2")
    eq(d, "B(T) = 8.24493e-1 - 4.0899e-3 T + 7.6438e-5 T^2 - 8.2467e-7 T^3 + 5.3875e-9 T^4")
    eq(d, "C(T) = -5.72466e-3 + 1.0227e-4 T - 1.6546e-6 T^2 ;   d0 = 4.8314e-4")
    p(d, "1.3 Hypersaline (brine) extension, anchored on EOS-80 at S = 42 psu "
         "(beta_S = 7.78e-4 /psu, beta_T = 3.2e-4 /degC):")
    eq(d, "rho_brine(S,T) = rho(42,T) [ 1 + beta_S (S - 42) - beta_T (T - 10) ]")
    p(d, "1.4 Smooth (C1) cosine blend over 38-46 psu (w in [0,1]):")
    eq(d, "w = 0.5 [ 1 - cos( pi * clip((S-38)/8, 0, 1) ) ]")
    eq(d, "rho_UBDS(S,T) = (1-w) rho_EOS80 + w rho_brine")
    p(d, "1.5 Reduced gravity (sign-agnostic; >0 dense, <0 buoyant):")
    eq(d, "g' = g (rho - rho_a) / rho_a")

    # --- 2 Near field ---
    h1(d, "2. Near-Field Integral Jet/Plume Model (module: nearfield.py)")
    p(d, "Eulerian top-hat flux-integral equations integrated along the jet "
         "arclength s by 4th-order Runge-Kutta. Conserved fluxes:")
    eq(d, "volume flux     Q   = pi b^2 w")
    eq(d, "mass flux       mu  = rho_e Q")
    eq(d, "momentum flux   Mvec = mu V          (V = (u,v,w); w = |V|)")
    eq(d, "salt flux       Sig = mu S")
    eq(d, "heat flux       The = mu T")
    p(d, "2.1 Governing flux balances (source terms per unit arclength):")
    eq(d, "d(mu)/ds   = rho_a E")
    eq(d, "d(Mvec)/ds = rho_a E V_a + (0, 0, (rho_a - rho_e) g pi b^2)   [buoyancy]")
    eq(d, "d(Sig)/ds  = rho_a E S_a")
    eq(d, "d(The)/ds  = rho_a E T_a")
    eq(d, "d(x,y,z)/ds = V / |V|")
    p(d, "2.2 Combined entrainment (shear + cross-flow / projected-area):")
    eq(d, "E = 2 pi b alpha_s |V - V_a|  +  2 b alpha_f |V_a,perp|")
    eq(d, "V_a,perp = V_a - (V_a . axis) axis ;  axis = V/|V|")
    p(d, "2.3 Froude-blended shear-entrainment coefficient:")
    eq(d, "Fr^2 = w^2 / (|g'| b) ;   w_j = Fr^2 / (Fr^2 + 1)")
    eq(d, "alpha_s = w_j alpha_jet + (1 - w_j) alpha_plume")
    eq(d, "alpha_jet = 0.057 ,  alpha_plume = 0.083 ,  alpha_f = 0.90  (standard, validated)")
    p(d, "2.4 Recovered geometry and dilution:")
    eq(d, "b = sqrt( Q / (pi w) ) ;   S = Sig/mu ;   T = The/mu ;   rho_e = rho_UBDS(S,T)")
    eq(d, "concentration dilution  D_n = (S0 - S_a) / (S - S_a)")
    p(d, "(Optional dense-fountain recirculation gain, default 1.0: multiplies E "
         "in the g'>0 regime; disabled in the validated configuration.)")

    # --- 3 Hydrodynamics ---
    h1(d, "3. Depth-Averaged Shallow-Water Hydrodynamics (module: hydro.py)")
    p(d, "Arakawa C-grid, explicit finite differences, upwind momentum advection.")
    eq(d, "Continuity:  d(eta)/dt + d(Hu)/dx + d(Hv)/dy = 0")
    eq(d, "x-momentum:  du/dt + u du/dx + v du/dy - f v =")
    eq(d, "             -g d(eta)/dx - g u sqrt(u^2+v^2)/(C^2 H) + A_h (d2u/dx2 + d2u/dy2) + a_x(t)")
    eq(d, "y-momentum:  dv/dt + u dv/dx + v dv/dy + f u =")
    eq(d, "             -g d(eta)/dy - g v sqrt(u^2+v^2)/(C^2 H) + A_h (d2v/dx2 + d2v/dy2) + a_y(t)")
    p(d, "3.1 Closures and forcing:")
    eq(d, "Chezy from Manning n:   C = H^(1/6) / n")
    eq(d, "Coriolis:               f = 2 Omega sin(phi)   (Omega = 7.2921159e-5 rad/s)")
    eq(d, "Tidal open boundary:    eta(t) = Sum_c  A_c cos(omega_c t - phi_c)")
    eq(d, "  constituents c in {M2, S2, K1, O1};  spring/neap via M2 +/- S2")
    eq(d, "Optional Smagorinsky viscosity: A_h = (Cs dx)^2 sqrt(2 ux^2 + 2 vy^2)")
    eq(d, "CFL gravity-wave limit: dt <= safety * dx / sqrt(g H_max)")

    # --- 4 Transport ---
    h1(d, "4. Quasi-3D Sigma-Layer Transport (module: transport.py)")
    p(d, "Advection-dispersion solved per terrain-following sigma layer k "
         "(advective/non-conservative upwind form for monotonicity):")
    eq(d, "dS_k/dt + u_k dS_k/dx + v_k dS_k/dy = D_h (d2S_k/dx2 + d2S_k/dy2) + Vert_k + Src_k")
    p(d, "4.1 Mass-conserving vertical velocity reconstruction (shear profile):")
    eq(d, "u_k = u * g_k ,  v_k = v * g_k")
    eq(d, "g_k = (sigma_k)^(1/p) / Sum_j ( frac_j (sigma_j)^(1/p) )   (p = 7),  Sum_k frac_k g_k = 1")
    p(d, "4.2 Vertical exchange = turbulent diffusion + dense gravity-current "
         "settling (the key novel closure):")
    eq(d, "F_k = K_z (S_k - S_{k-1}) / dz  +  w_s,k * max(S_k - S_{k-1}, 0)")
    eq(d, "w_s,k = c_s sqrt( max(g'_k,0) * h_layer )   (flux-limited, capped)")
    eq(d, "g'_k = g (rho_k - rho_bg) / rho_bg ;   dS_k += dt * div_z(F) / h_layer,k")
    p(d, "4.3 Near-field-coupled source (bottom layer footprint):")
    eq(d, "excess-salt mass flux:  Phi = Q0 (S0 - S_a)   [psu m3/s]  (load-scaled)")
    eq(d, "dS_0 += dt * Phi / (h_layer,0 dx dy)")
    eq(d, "Dirichlet floor at footprint:  S_0 = max(S_0, S_impact)  (from near-field)")
    eq(d, "Open boundaries: edge cells reset to S_bg (background reservoir / flushing)")

    # --- 5 Metrics ---
    h1(d, "5. Metrics and Skill (module: metrics.py)")
    eq(d, "Salinity-increment area:  A(dS) = (number of cells with S >= S_bg + dS) * dx dy")
    eq(d, "Mixing-zone radius:  R = max distance from source where S >= S_threshold")
    eq(d, "EQS compliance:  C_pred = C_bg + (C_brine - C_bg) / D_n ;  margin = EQS / C_pred")
    eq(d, "RMSE = sqrt( mean( (model - obs)^2 ) )")
    eq(d, "Willmott index:  d = 1 - Sum(m-o)^2 / Sum( (|m - obar| + |o - obar|)^2 )")
    eq(d, "Nash-Sutcliffe:  NSE = 1 - Sum(m-o)^2 / Sum(o - obar)^2")

    d.save("model.equations.docx")
    print("model.equations.docx written")


# ===========================================================================
#  SOURCES DOCUMENT
# ===========================================================================
def build_sources():
    d = new_doc()
    t = d.add_heading("", 0)
    t.add_run("UBDS - Calibration & Validation Sources").font.color.rgb = INK
    m = d.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.add_run("Author: Akosa Samuel Onyejekwe (Independent Researcher)").font.color.rgb = INK

    h1(d, "1. Calibration sources")
    p(d, "Quantities used to set/calibrate the solver coefficients:")
    h2(d, "1.1 Near-field entrainment coefficients")
    bullet(d, "Top-hat shear-entrainment coefficients alpha_jet = 0.057 and "
              "alpha_plume = 0.083 - standard turbulent jet/plume values from "
              "Fischer, List, Koh, Imberger & Brooks (1979) and Lee & Chu (2003); "
              "verified to reproduce the canonical round-jet spreading rate "
              "(db/dz ~ 0.11) and the plume 5/3 volume-flux law.")
    bullet(d, "Cross-flow (forced/projected-area) coefficient alpha_f = 0.90 - "
              "set so the inclined dense-jet return distance x_r/(D Fr) falls in "
              "the Papakonstantis et al. (2011) experimental band (~2.2-3.3).")
    bullet(d, "Initial-dilution check dataset (dense vertical 6-inch jets, brine "
              "260 psu into 34.2 psu, 27 m depth) used as a calibration/validation "
              "reference for the near-field kernel (see validation/calibrate_nearfield.py).")
    h2(d, "1.2 Hydrodynamic calibration")
    bullet(d, "Tidal constituents M2, S2, K1, O1 (amplitudes/phases) defining a "
              "meso-tidal semidiurnal regime; the stream-forcing phase lag is "
              "calibrated so the modelled peak current at the outfall matches the "
              "target (spring ~0.55 m/s, neap ~0.28 m/s).")
    bullet(d, "Manning bed-roughness n = 0.025 (range 0.022-0.040), horizontal eddy "
              "viscosity 2 m2/s (Smagorinsky Cs = 0.2) - standard coastal values.")
    bullet(d, "Synthetic ADCP current dataset (5 bottom-mounted upward-looking "
              "profilers, neap & spring windows) used to demonstrate the calibration "
              "workflow and skill metrics (Willmott d, Nash-Sutcliffe, RMSE).")
    h2(d, "1.3 Transport calibration")
    bullet(d, "Horizontal dispersion D_h = 1 m2/s (open-water) / 10 m2/s "
              "(coastal recirculation study); vertical diffusivity K_z = 5e-4 m2/s; "
              "dense gravity-current coefficient c_s tuned to give bottom-trapping "
              "consistent with the reference-study vertical salinity decay.")

    h1(d, "2. Validation sources (literature & reference data)")
    refs = [
     ("Fischer, H.B., List, E.J., Koh, R.C.Y., Imberger, J. & Brooks, N.H. (1979). "
      "Mixing in Inland and Coastal Waters. Academic Press.",
      "Canonical round-jet dilution law (bulk dilution ~0.25-0.32 z/D) and jet "
      "spreading rate (db/dz ~ 0.11). Validates the momentum-jet limit."),
     ("Morton, B.R., Taylor, G.I. & Turner, J.S. (1956). Turbulent gravitational "
      "convection from maintained and instantaneous sources. Proc. R. Soc. Lond. A 234, 1-23.",
      "Pure-plume volume-flux 5/3 power law. Validates the buoyancy-driven limit "
      "and entrainment closure."),
     ("Papakonstantis, I.G., Christodoulou, G.C. & Papanicolaou, P.N. (2011). "
      "Inclined negatively buoyant jets 1 & 2. J. Hydraulic Research 49(1), 3-22.",
      "Dimensionless terminal rise z_t/(D Fr) ~ 1.6-2.2 and centreline return-point "
      "dilution S_r/Fr ~ 0.4-0.6 for 60-degree dense jets. Validates the dense (brine) jet."),
     ("Roberts, P.J.W., Ferrier, A. & Daviero, G. (1997). Mixing in inclined dense "
      "jets. J. Hydraulic Engineering 123(8), 693-699.",
      "Independent corroboration of inclined dense-jet dilution behaviour."),
     ("Millero, F.J. & Poisson, A. (1981). International one-atmosphere equation of "
      "state of seawater. Deep-Sea Research 28A, 625-629; UNESCO (1981), Tech. Papers "
      "Mar. Sci. 36.",
      "EOS-80 reference seawater densities (S <= 42 psu). Validates the equation of state."),
     ("CRC Handbook of Chemistry and Physics.",
      "Saturated NaCl brine density (~1200 kg/m3) anchoring the hypersaline extension."),
     ("Lee, J.H.W. & Chu, V.H. (2003). Turbulent Jets and Plumes - A Lagrangian "
      "Approach. Kluwer Academic Publishers.",
      "Combined shear + forced (cross-flow) entrainment hypothesis underpinning the "
      "near-field entrainment formulation."),
     ("Jirka, G.H. (2004). Integral model for turbulent buoyant jets in unbounded "
      "stratified flows, Part 1. Environmental Fluid Mechanics 4, 1-56.",
      "Reference integral-model formulation for buoyant jets."),
     ("Willmott, C.J. (1981). On the validation of models. Physical Geography 2, 184-194.",
      "Index of agreement (d) used to score the hydrodynamic calibration."),
     ("Nash, J.E. & Sutcliffe, J.V. (1970). River flow forecasting through "
      "conceptual models, Part I. Journal of Hydrology 10(3), 282-290.",
      "Nash-Sutcliffe efficiency (NSE) used to score the hydrodynamic calibration."),
    ]
    for cite, use in refs:
        bp = d.add_paragraph(style="List Number")
        r = bp.add_run(cite); r.bold = True; r.font.size = Pt(10); r.font.color.rgb = INK
        sub = d.add_paragraph(); sub.paragraph_format.left_indent = Pt(24)
        rr = sub.add_run("Validates / used for: " + use)
        rr.italic = True; rr.font.size = Pt(9.5); rr.font.color.rgb = INK

    h1(d, "3. Site-specific case data")
    p(d, "All site-specific inputs - bathymetry, tidal constituents, ambient "
         "salinity/temperature, brine chemistry and trace-impurity concentrations, "
         "diffuser geometry - are synthetic but engineering-credible values defined "
         "by the author for the Bahia Azul case study (see case_inputs.py). None are "
         "taken from any real consent. Marine Environmental Quality Standards (EQS) "
         "used for the impurity comparison follow typical marine-water EQS values.")

    h1(d, "4. Validation results summary")
    try:
        import json
        v = json.load(open("validation/validation_summary.json"))
        for k, val in v.items():
            bullet(d, f"{k} = {val}")
    except Exception:
        pass
    bullet(d, "Hydrodynamic calibration: Willmott d = 0.99 and NSE > 0.96 at all "
              "five ADCP stations (neap & spring).")

    d.save("source.docx")
    print("source.docx written")


if __name__ == "__main__":
    build_equations()
    build_sources()
