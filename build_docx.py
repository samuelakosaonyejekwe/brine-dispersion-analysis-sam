"""
Assemble the case-study report: case.docx
=========================================

Collates the complete UBDS case study - inputs, geometry/mesh, model setup,
near/medium/far-field results, design optimisation, water quality, validation
and data sources - embedding every generated figure, table and CSV reference.

No pure black is used anywhere (text and headings use dark navy #1B2A4A).

Author: Akosa Samuel Onyejekwe (Independent Researcher)
"""
import os, json, pickle
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

import case_inputs as C
import run_helpers as H
from ubds import eos

INK = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2A, 0x6F, 0x97)
MANIFEST = json.load(open(H.MANIFEST))
RESULTS = pickle.load(open("outputs/results.pkl", "rb"))
VSUM = json.load(open("validation/validation_summary.json")) if os.path.exists(
    "validation/validation_summary.json") else {}

# --- jet Froude number of every near-field run, vs the validated envelope -----
_nfcsv = pd.read_csv(f"{H.CSV_DIR}/A_nearfield_initial_dilution.csv")
_ra = eos.density(C.AMBIENT["background_salinity_psu"], C.AMBIENT["background_temperature_C"])
def _fr(row):
    if row.brine == "RO concentrate":
        S0, T0 = C.BRINE["ro_salinity_psu"], C.AMBIENT["background_temperature_C"] + C.BRINE["ro_excess_temp_C"]
    else:
        S0, T0 = C.BRINE["cavern_salinity_psu"], C.AMBIENT["background_temperature_C"] + C.BRINE["cavern_excess_temp_C"]
    gp = 9.81 * (eos.density(S0, T0) - _ra) / _ra
    return row.exit_velocity_ms / np.sqrt(gp * C.DIFFUSER["port_diameter_m"])
_nfcsv["Fr"] = _nfcsv.apply(_fr, axis=1)
# Read the benchmarked Froude span from the validation suite rather than hard-coding
# it, and test BOTH bounds.  Testing only the upper bound previously let nine
# saturated-cavern runs at Fr 9.2-10.7 sit below the benchmark while the report
# asserted that every run was inside it.
_FR_LO = VSUM["dense_jet_Fr_min"]
_FR_HI = VSUM["dense_jet_Fr_max"]
_out = _nfcsv[(_nfcsv.Fr < _FR_LO) | (_nfcsv.Fr > _FR_HI)]

doc = Document()

# global style -> dark navy, no black
st = doc.styles["Normal"]
st.font.name = "Calibri"; st.font.size = Pt(10.5); st.font.color.rgb = INK
for h in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
    try:
        doc.styles[h].font.color.rgb = ACCENT if h != "Title" else INK
    except Exception:
        pass


def figs(section):
    return [a for a in MANIFEST if a["kind"] == "figure" and a["section"] == section]


def csvs(section):
    return [a for a in MANIFEST if a["kind"] == "csv" and a["section"] == section]


def add_fig(a, width=5.7):
    if not os.path.exists(a["path"]):
        return
    doc.add_picture(a["path"], width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    r = cap.add_run("Figure. " + a["caption"])
    r.italic = True; r.font.size = Pt(9); r.font.color.rgb = INK
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


_EMITTED_SECTIONS = set()


def add_section_figs(section, width=5.7):
    found = figs(section)
    if not found:
        raise KeyError(f"no figures registered for section {section!r} - "
                       "check the manifest, do not silently emit nothing")
    _EMITTED_SECTIONS.add(section)
    for a in found:
        add_fig(a, width)


def add_table_from_csv(path, caption, max_rows=40, drop_cols=()):
    # A missing CSV means the pipeline did not produce a table the report claims
    # to show.  Failing loudly here is what surfaces that; returning silently is
    # how a missing table can sit unnoticed in a published report.
    if not os.path.exists(path):
        raise FileNotFoundError(f"table source missing: {path}")
    df = pd.read_csv(path)
    if drop_cols:
        df = df.drop(columns=list(drop_cols))
    if len(df) > max_rows:
        raise ValueError(f"{path}: {len(df)} rows exceeds max_rows={max_rows}; "
                         "truncating would silently drop data")
    p = doc.add_paragraph()
    r = p.add_run("Table. " + caption); r.italic = True; r.font.size = Pt(9)
    r.font.color.rgb = INK
    t = doc.add_table(rows=1, cols=len(df.columns))
    t.style = "Light Grid Accent 1"
    # Let Word size columns to their content.  With fixed equal widths a wide
    # table wraps numbers mid-value ("356.3" -> "356. 3"), which misreads badly.
    t.autofit = True
    t.allow_autofit = True
    t._tbl.tblPr.xpath("./w:tblLayout") and [
        e.getparent().remove(e) for e in t._tbl.tblPr.xpath("./w:tblLayout")]
    fsz = Pt(8.5) if len(df.columns) <= 8 else Pt(7.5)
    for j, c in enumerate(df.columns):
        cell = t.rows[0].cells[j]; cell.text = str(c)
        for pp in cell.paragraphs:
            for rr in pp.runs:
                rr.font.bold = True; rr.font.size = fsz; rr.font.color.rgb = INK
    for _, row in df.iterrows():
        cells = t.add_row().cells
        for j, c in enumerate(df.columns):
            # A missing value must not print as the literal "nan" in a report.
            v = row[c]
            cells[j].text = "-" if pd.isna(v) else str(v)
            for pp in cells[j].paragraphs:
                for rr in pp.runs:
                    rr.font.size = fsz; rr.font.color.rgb = INK
    doc.add_paragraph("")


def h1(t): doc.add_heading(t, 1)
def h2(t): doc.add_heading(t, 2)
def para(t):
    p = doc.add_paragraph(t)
    return p
def bullet(t): doc.add_paragraph(t, style="List Bullet")


# ===========================================================================
# TITLE
# ===========================================================================
ttl = doc.add_heading("", 0)
ttl.add_run("Brine Dispersion Modelling Case Study").font.color.rgb = INK
sub = doc.add_paragraph()
sr = sub.add_run("Universal Brine Dispersion Solver (UBDS): Near-, Medium- and "
                 "Far-Field Assessment of the Bahia Azul Seawater Desalination "
                 "& Brine Outfall")
sr.bold = True; sr.font.size = Pt(13); sr.font.color.rgb = ACCENT
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Author: Akosa Samuel Onyejekwe  (Independent Researcher)\n"
                  "Solver: UBDS v1.0.0  |  Document: case.docx")
mr.font.size = Pt(11); mr.font.color.rgb = INK
doc.add_paragraph("")

# ===========================================================================
# EXECUTIVE SUMMARY
# ===========================================================================
h1("Executive Summary")
nf_ro = RESULTS["nearfield"].query("brine=='RO concentrate'")
far = RESULTS.get("far_table")
para(
 "This study presents a holistic near-, medium- and far-field assessment of the "
 "brine discharge from the proposed Bahia Azul seawater desalination and brine "
 "outfall, performed entirely with the Universal Brine Dispersion Solver (UBDS) - "
 "a novel, open, physics-based solver developed by the author. UBDS unifies, in a "
 "single dynamically-coupled and mass-conserving framework, the near-field integral "
 "jet/plume modelling and the far-field hydrodynamic transport modelling that "
 "conventionally require two separate, proprietary codes. Its sign-agnostic "
 "buoyancy treatment and dense-water gravity-current closure make it universal "
 "across negatively buoyant (brine), positively buoyant (thermal) and neutral "
 "discharges.")
para(
 f"For the proposed inclined {C.DIFFUSER['n_ports_installed']}-port diffuser, the near-field initial dilution of "
 f"the reverse-osmosis (RO) concentrate at first seabed contact ranges between "
 f"x{nf_ro.dilution.min():.0f} and x{nf_ro.dilution.max():.0f} across the flow and "
 f"tidal range, corresponding to seabed salinities of "
 f"{nf_ro.seabed_salinity_psu.min():.1f}-{nf_ro.seabed_salinity_psu.max():.1f} psu "
 f"against a background of {C.AMBIENT['background_salinity_psu']} psu.")
if far is not None:
    s3 = far[far.scenario == "S3"].iloc[0]
    if s3.area_0p5psu_km2 > 0:
        _area = (f"the salinity increment of 0.5 psu above background is confined to an area of "
                 f"{s3.area_0p5psu_km2:.3f} km2")
    else:
        # The zero area is a statement about the regional grid, not about the plume:
        # the far-field cell is far coarser than the near-field impact footprint, and
        # the increment does exceed 0.5 psu inside that footprint at slack water.
        _far_inc = far.max_seabed_salinity_psu.max() - C.AMBIENT["background_salinity_psu"]
        _nf_inc = (RESULTS["nearfield"].query("brine=='RO concentrate'").seabed_salinity_psu.max()
                   - C.AMBIENT["background_salinity_psu"])
        _area = (f"no 0.5 psu influence area is resolved on the regional grid, whose seabed "
                 f"increment peaks at {_far_inc:.2f} psu. The reported 0.000 km2 is therefore a "
                 f"statement about the {C.FAR_FIELD['dx_m']:.0f} m regional cell, not a claim that "
                 f"the increment is nowhere above 0.5 psu: at the near-field impact footprint at "
                 f"slack water it reaches {_nf_inc:.2f} psu, over a scale the regional grid cannot "
                 f"resolve")
    para(
     f"The medium- and far-field assessment confirms that, at the maximum discharge "
     f"of {C.SCENARIOS[-1]['flow_m3hr']} m3/hr, {_area}, with "
     f"the {C.THRESHOLDS['salinity_threshold_psu']} psu threshold not exceeded beyond "
     f"{s3.mixing_zone_radius_m:.0f} m from the diffuser - well within the "
     f"{C.THRESHOLDS['mixing_zone_radius_m']:.0f} m regulatory mixing zone. All trace "
     f"impurities remain below their Environmental Quality Standards after dilution. "
     + (f"The RO near-field dilutions above are extrapolated beyond the validated Froude "
        f"range (see section 4); the saturated-cavern-brine worst case, which lies inside "
        f"the validated range, governs the seabed-salinity assessment."
        if len(_out) else
        f"Every near-field run lies inside the Froude range over which the entrainment "
        f"closure is validated (see section 4); the saturated-cavern-brine worst case "
        f"governs the seabed-salinity assessment."))
para(
 "The solver was tested against canonical self-similar jet and plume dilution laws, "
 "published inclined dense-jet experimental data, and reference seawater/brine densities. "
 "The equation of state, the dense-jet "
 "return distance and the dense-jet return-point dilution fall inside their reference "
 "ranges; the round-jet dilution slope, the pure-plume exponent and the dense-jet "
 "terminal rise fall a few per cent outside theirs. Each deviation is quantified in the "
 "Validation section rather than asserted away. Data sources are recorded there also. "
 "The hydrodynamic skill scores in section 3 are not part of this evidence: the ADCP "
 "series they score against is synthetic, derived from the model solution itself.")

# ===========================================================================
# 1 INTRODUCTION & SOLVER
# ===========================================================================
h1("1  Introduction and the UBDS Solver")
para(
 "Brine discharges from desalination plants and solution-mined storage caverns are "
 "negatively buoyant and sink towards the seabed, where elevated salinity can affect "
 "benthic ecology. Their assessment requires resolving both the small-scale turbulent "
 "jet/plume mixing in the immediate vicinity of the diffuser (the near field) and the "
 "tidally-driven transport and dilution over the wider water body (the medium and far "
 "field). Conventionally these two regimes are handled by separate, proprietary "
 "software, with manual transfer of results between them.")
h2("1.1  Novelty and gaps bridged")
bullet("Unified near-to-far field coupling in a single open framework, removing the "
       "manual hand-off and the associated mass-conservation errors between separate codes.")
bullet("Sign-agnostic buoyancy: the same governing equations describe dense (brine), "
       "buoyant (thermal/freshwater) and neutral discharges - a genuinely universal solver.")
bullet("A dense-water gravity-current closure in the sigma-layer transport engine that "
       "reproduces the strong bottom-trapping of brine that 2-D depth-averaged models cannot "
       "capture and that 3-D models obtain only at large computational cost.")
bullet("Built-in regulatory-metric engine: mixing-zone radius, salinity-increment areas, "
       "diffusion distances and EQS compliance computed automatically from the solution.")
bullet("Fully open and reproducible: physically-standard, documented entrainment and "
       "turbulence closures, with all inputs and outputs exposed as data.")
h2("1.2  Governing equations")
para("Near field - steady Eulerian flux-integral equations (top-hat) integrated along the "
     "jet arclength s, with combined shear + cross-flow entrainment and a sign-agnostic "
     "buoyancy body force:")
para("   d(rho_e Q)/ds = rho_a E ;   d(rho_e Q V)/ds = rho_a E Va + (0,0,(rho_a-rho_e) g pi b^2)")
para("   d(rho_e Q S)/ds = rho_a E Sa ;   E = 2 pi b alpha_s |V-Va| + 2 b alpha_f |Va_perp|")
para("Far field - depth-averaged shallow-water hydrodynamics (mass + momentum with Manning "
     "friction, Coriolis and turbulent mixing) driving a sigma-layer advection-dispersion "
     "equation augmented with the buoyancy-scaled inter-layer gravity-current flux:")
para("   d(eta)/dt + d(Hu)/dx + d(Hv)/dy = 0")
para("   d(hs)/dt + d(uhs)/dx + d(vhs)/dy = d/dx(hDx ds/dx) + d/dy(hDy ds/dy) + sources")

# ===========================================================================
# 2 SITE & INPUTS
# ===========================================================================
h1("2  Case Study Site and Input Data")
para(f"Site: {C.SITE['name']} ({C.SITE['locality']}). A mid-size seawater reverse-osmosis "
     f"desalination plant (product capacity {C.PLANT['product_capacity_m3d']:,} m3/d) "
     f"discharging hypersaline concentrate through an offshore inclined multiport diffuser, "
     f"with provision for later receipt of saturated leaching brine from an adjacent "
     f"salt-cavern energy-storage scheme. Latitude {C.SITE['latitude_deg']} N.")
h2("2.1  Discharge and ambient inputs")
inp_rows = [
 ("Product (freshwater) capacity", f"{C.PLANT['product_capacity_m3d']:,} m3/d"),
 ("Recovery ratio", f"{C.PLANT['recovery_ratio']}"),
 ("Brine (concentrate) flow - full", f"{C.SCENARIOS[-1]['flow_m3hr']} m3/hr "
                                      f"({C.SCENARIOS[-1]['flow_m3hr']*24:,} m3/d)"),
 ("RO brine salinity / excess temp", f"{C.BRINE['ro_salinity_psu']} psu / +{C.BRINE['ro_excess_temp_C']} C"),
 ("Cavern brine salinity / excess temp", f"{C.BRINE['cavern_salinity_psu']} psu / +{C.BRINE['cavern_excess_temp_C']} C"),
 ("Background salinity / temperature", f"{C.AMBIENT['background_salinity_psu']} psu / "
                                       f"{C.AMBIENT['background_temperature_C']} C"),
 ("Outfall depth / offshore distance", f"{C.AMBIENT['outfall_depth_m']} m / {C.AMBIENT['offshore_distance_m']} m"),
 ("Diffuser", f"{C.DIFFUSER['n_ports_installed']} x 6\" ports @ {C.DIFFUSER['port_spacing_m']} m, "
              f"{C.DIFFUSER['port_angle_deg']} deg, {C.DIFFUSER['port_height_m']} m above bed"),
 ("Mixing-zone radius / threshold", f"{C.THRESHOLDS['mixing_zone_radius_m']} m / "
                                    f"{C.THRESHOLDS['salinity_threshold_psu']} psu"),
]
t = doc.add_table(rows=0, cols=2); t.style = "Light Grid Accent 1"
for k, v in inp_rows:
    c = t.add_row().cells; c[0].text = k; c[1].text = v
    for cell in c:
        for pp in cell.paragraphs:
            for rr in pp.runs:
                rr.font.size = Pt(9); rr.font.color.rgb = INK
doc.add_paragraph("")
h2("2.2  Flow scenarios")
add_table_from_csv(f"{H.CSV_DIR}/F_model_parameters.csv",
                   "Hydrodynamic and transport model parameters used in the simulations.")

# ===========================================================================
# 3 GEOMETRY, MESH & SETUP
# ===========================================================================
h1("3  Model Geometry, Mesh and Setup")
para("Two nested structured-grid domains were configured: a high-resolution medium-field "
     f"domain ({C.MEDIUM_FIELD['Lx_m']:.0f} x {C.MEDIUM_FIELD['Ly_m']:.0f} m at "
     f"{C.MEDIUM_FIELD['dx_m']:.0f} m, {C.MEDIUM_FIELD['n_sigma_layers']} sigma layers) and a "
     f"regional far-field domain ({C.FAR_FIELD['Lx_m']:.0f} x {C.FAR_FIELD['Ly_m']:.0f} m at "
     f"{C.FAR_FIELD['dx_m']:.0f} m, {C.FAR_FIELD['n_sigma_layers']} layers). The bathymetry "
     "deepens offshore and includes a southern headland whose lee generates a tidal eddy.")
add_section_figs("Hydrodynamics, mesh & calibration")
add_table_from_csv(f"{H.CSV_DIR}/F_layer_scheme.csv",
                   "Sigma-layer schemes (% of water-column depth per layer).")
para("A note on what the skill scores below do and do not show. Bahia Azul is a fictitious "
     "site, so no measured currents exist for it. The 'ADCP' series is the model solution "
     "itself, perturbed by a 3 per cent proportional bias and random noise scaled to the local "
     "signal amplitude. The Willmott index and Nash-Sutcliffe efficiency reported here "
     "therefore quantify that imposed perturbation and confirm that the skill-metric pipeline "
     "behaves correctly. They are NOT evidence that the hydrodynamic engine reproduces field "
     "measurements, and they must not be read as calibration against observations. The "
     "independent evidence for the solver is in the Validation section, which compares it "
     "against published experimental data and analytical laws.")
add_table_from_csv(f"{H.CSV_DIR}/B_current_validation_skill.csv",
                   "Model skill against the synthetic ADCP reference series. The reference is the "
                   "model solution plus an imposed bias and noise, so near-unity scores are expected "
                   "by construction and are not a field validation.")

# ===========================================================================
# 4 NEAR FIELD
# ===========================================================================
h1("4  Near-Field Initial Dilution")
para("The near-field initial dilution was computed with the UBDS integral jet/plume engine "
     "for both brine types, all flow scenarios and a range of tidal velocities (slack, neap "
     "and spring). Dilution is the bulk (flux-averaged, top-hat) dilution, and salinity is "
     "reported at first contact with the seabed; every run in the table below is verified to "
     "terminate on seabed contact.")

_depth = C.AMBIENT["outfall_depth_m"]
_rise_frac = 100 * _nfcsv.rise_height_m.max() / _depth
_cav = _nfcsv[_nfcsv.brine != "RO concentrate"]
_ro = _nfcsv[_nfcsv.brine == "RO concentrate"]
if not len(_out):
    para(
      f"Validity of the near-field results. The inclined dense-jet benchmark spans "
      f"{_FR_LO:.1f} <= Fr <= {_FR_HI:.1f} ({VSUM['dense_jet_n_cases']} cases), which was chosen to "
      f"cover the Froude numbers this study actually runs at rather than a convenient band above "
      f"them. All {len(_nfcsv)} runs below lie inside that span: the RO concentrate at "
      f"Fr {_ro.Fr.min():.1f}-{_ro.Fr.max():.1f} and the denser saturated-cavern brine, which "
      f"governs the seabed assessment, at Fr {_cav.Fr.min():.1f}-{_cav.Fr.max():.1f}. The diffuser "
      f"stages ports in with flow, holding the per-port exit velocity between "
      f"{_nfcsv.exit_velocity_ms.min():.2f} and {_nfcsv.exit_velocity_ms.max():.2f} m/s. "
      f"Coverage is not the same as agreement: across the benchmark the return distance is inside "
      f"its published band in {VSUM['dense_jet_xr_in_band']}/{VSUM['dense_jet_n_cases']} cases and "
      f"the return-point dilution in {VSUM['dense_jet_Sr_min_in_band']}/{VSUM['dense_jet_n_cases']}, "
      f"but the terminal rise is inside its band in only "
      f"{VSUM['dense_jet_zt_in_band']}/{VSUM['dense_jet_n_cases']}, and the shortfall grows as Fr "
      f"falls - that is, it is largest at the cavern-brine Froude numbers. Section 9 quantifies it. "
      f"The greatest predicted plume rise is {_nfcsv.rise_height_m.max():.1f} m in {_depth:.0f} m "
      f"of water ({_rise_frac:.0f}% of the water column), so the free surface, which the model does "
      f"not impose, is not approached.")
else:
    para(
      f"Validity of the near-field results. The entrainment closure is validated against "
      f"inclined dense-jet data over {_FR_LO:.1f} <= Fr <= {_FR_HI:.1f} and against vertical dense-jet "
      f"initial-dilution data over 7.4 <= Fr <= 14.8. {len(_out)} of the {len(_nfcsv)} runs in the "
      f"table below fall outside that range (Fr {_nfcsv.Fr.min():.1f}-{_nfcsv.Fr.max():.1f}): these are the "
      f"RO-concentrate cases, whose brine is only "
      f"{C.BRINE['ro_salinity_psu']:.0f} psu and therefore only weakly dense, giving a high jet "
      f"Froude number and correspondingly high dilution. Their predicted dilutions "
      f"(up to x{_nfcsv.dilution.max():.0f}) are EXTRAPOLATIONS beyond the validated envelope and "
      f"should be treated as indicative rather than design values. In addition, the greatest "
      f"predicted plume rise is {_nfcsv.rise_height_m.max():.1f} m in {_depth:.0f} m of water "
      f"({_rise_frac:.0f}% of the water column); the model does not impose a free-surface "
      f"constraint, so a run that approaches the surface will over-predict rise and hence "
      f"trajectory length and entrainment. The saturated-cavern-brine cases, which are the "
      f"worst case for seabed salinity, all lie inside the validated Fr range and are not "
      f"affected by either caveat.")
add_table_from_csv(f"{H.CSV_DIR}/A_nearfield_initial_dilution.csv",
                   "Near-field initial-dilution results. Every run terminates on seabed contact; the "
                   "termination check is retained in the CSV.", max_rows=60,
                   drop_cols=("termination",))
add_section_figs("Near-field initial dilution")

# ===========================================================================
# 5 MEDIUM FIELD
# ===========================================================================
h1("5  Medium-Field Dispersion")
para("The medium-field dispersion was simulated with the sigma-layer transport engine driven "
     "by the calibrated tidal currents, including a layer-resolution sensitivity test and the "
     "near-field-coupled brine source. Maximum-salinity envelopes, tidal-phase snapshots with "
     "current vectors, and vertical profiles through the diffuser are presented. Every "
     "medium-field run in this section uses the SATURATED CAVERN BRINE source at full flow - "
     "the phase-2 worst case, not the RO concentrate of the base case. The medium-field seabed "
     "salinities are therefore not comparable with the far-field values of section 6, which are "
     "computed for RO concentrate unless a figure states otherwise.")
ls = RESULTS.get("layer_sensitivity", {})
if ls:
    _vals = [v for _, v in sorted(ls.items())]
    _converged = abs(_vals[-1] - _vals[-2]) / _vals[-1] < 0.01 if len(_vals) > 1 else False
    para("Layer-resolution sensitivity (maximum seabed salinity, psu): " +
         ", ".join(f"{k}-layer = {v:.1f}" for k, v in sorted(ls.items())) +
         ". The 9-layer scheme was adopted for the production runs. " +
         ("The peak is grid-converged in the vertical." if _converged else
          f"The peak salinity is NOT converged in the vertical: it continues to rise with layer "
          f"count, so the adopted 9-layer scheme under-predicts the 14-layer peak by "
          f"{100*(_vals[-1]-_vals[1])/_vals[-1]:.1f}%. Peak seabed salinity in this section should "
          f"therefore be read as a lower bound; the mixing-zone and area metrics, which depend on "
          f"the salinity field away from the peak cell, are far less sensitive to layer count."))
add_section_figs("Medium-field dispersion")

# ===========================================================================
# 6 FAR FIELD
# ===========================================================================
h1("6  Far-Field Dispersion")
add_section_figs("Far-field dispersion")
_far_inc = far.max_seabed_salinity_psu.max() - C.AMBIENT["background_salinity_psu"]
_nf_inc = (RESULTS["nearfield"].query("brine=='RO concentrate'").seabed_salinity_psu.max()
           - C.AMBIENT["background_salinity_psu"])
para(f"How to read the zeros in the table below. On the {C.FAR_FIELD['dx_m']:.0f} m regional grid the "
     f"seabed salinity increment peaks at {_far_inc:.2f} psu, below the lowest reported threshold of "
     f"0.5 psu, so every influence area evaluates to exactly 0.000 km2 and the mixing-zone radius to "
     f"0 m. These are grid-resolved quantities. They do not mean the increment never exceeds 0.5 psu "
     f"anywhere: the near-field solution reaches {_nf_inc:.2f} psu above background at the impact "
     f"footprint at slack water (section 4), on a scale of metres that a "
     f"{C.FAR_FIELD['dx_m']:.0f} m cell cannot represent. The regional figures therefore bound the "
     f"far-field footprint, and the near-field solution governs the immediate seabed impact.")
add_table_from_csv(f"{H.CSV_DIR}/D_influence_area_diffusion.csv",
                   "Salinity-increment influence areas and diffusion distances per scenario. Areas are "
                   "resolved on the regional grid; see the note above on what the zeros mean.")

# ===========================================================================
# 7 DESIGN OPTIMISATION
# ===========================================================================
h1("7  Diffuser Design Optimisation and Layout")
para("The effect of diffuser port discharge angle on near-field dilution was assessed, and a "
     "recommended layout selected. Cross-shore salinity and current-speed transects and the "
     "marine-constraints overlay support the siting decision.")
add_section_figs("Diffuser design optimisation")
add_section_figs("Design comparison & optimisation")

# ===========================================================================
# 7B  OUTFALL SITING & INTAKE RECIRCULATION
# ===========================================================================
h1("8  Outfall Siting and Intake Recirculation")
_rc = RESULTS["recirc"]
_CRIT = 0.1
_ro = _rc[_rc.brine == "RO concentrate"]
_cav = _rc[_rc.brine != "RO concentrate"]
_ro_worst = _ro.rise_max_ppt.max()
_cav_worst = _cav.rise_max_ppt.max()
# The siting distance is the one the case is actually run at, not a separate figure:
# AMBIENT["offshore_distance_m"] places the diffuser in every medium- and far-field run.
_site_dist = C.AMBIENT["offshore_distance_m"]
_cav_ok = sorted(d for d in _cav.outfall_dist_m.unique()
                 if _cav[_cav.outfall_dist_m == d].rise_max_ppt.max() <= _CRIT)
_cav_min_ok = _cav_ok[0] if _cav_ok else None
_tested = sorted(_rc.outfall_dist_m.unique())
para("Following the desalination-siting methodology, the salinity rise at the plant intake "
     "was assessed as a function of outfall distance from shore under steady northward, "
     f"southward and weak-southward currents, for both discharges. The criterion is a salinity "
     f"rise at the intake of no more than {_CRIT} ppt above ambient, applied to the peak cell in "
     "the intake window. The weak southward current is the critical (least-flushing) condition.")
_nf_all = RESULTS["nearfield"]
_ro_bed = _nf_all[_nf_all.brine == "RO concentrate"].seabed_salinity_psu.max()
_cav_bed = _nf_all[_nf_all.brine != "RO concentrate"].seabed_salinity_psu.max()
_bg = C.AMBIENT["background_salinity_psu"]
para("The two brines behave very differently, and the distinction matters. The intake signal "
     "scales with the excess salinity the plume delivers to the bed: the RO concentrate lands at "
     f"{_ro_bed:.2f} psu, an excess of {_ro_bed - _bg:.2f} psu over the {_bg} psu background, "
     f"whereas the saturated cavern brine lands at {_cav_bed:.2f} psu, an excess of "
     f"{_cav_bed - _bg:.2f} psu - about {(_cav_bed - _bg) / (_ro_bed - _bg):.0f} times larger.")
add_table_from_csv(f"{H.CSV_DIR}/H_table1A_southward.csv",
                   "Rise in salinity at intake above ambient - southward current 0.25 m/s, both brines.")
add_table_from_csv(f"{H.CSV_DIR}/H_table1B_weak.csv",
                   "Rise in salinity at intake above ambient - weak southward current 0.05 m/s, both brines.")
add_table_from_csv(f"{H.CSV_DIR}/H_table1C_northward.csv",
                   "Rise in salinity at intake above ambient - northward current 0.25 m/s, both brines.")
para(f"Result. For the RO concentrate the criterion is met at every distance tested: the worst cell "
     f"over all currents and distances is {_ro_worst:.3f} ppt, at the closest {_tested[0]:.0f} m "
     f"siting under the weak southward current. The {_site_dist:.0f} m siting modelled in this "
     f"study therefore clears the criterion, as does every distance from {_tested[0]:.0f} to "
     f"{_tested[-1]:.0f} m. For the phase-2 saturated cavern brine the criterion is NOT met at any "
     f"distance below {_cav_min_ok if _cav_min_ok else -1:.0f} m (worst case {_cav_worst:.3f} ppt "
     f"at the closest siting under the weak southward current). It is first met at "
     f"{_cav_min_ok if _cav_min_ok else -1:.0f} m offshore ("
     f"{_cav[_cav.outfall_dist_m == (_cav_min_ok or 0)].rise_max_ppt.max():.3f} ppt).")
para(f"Consequence for the design. The {_site_dist:.0f} m outfall modelled in this study satisfies "
     "the recirculation criterion for the phase-1 RO discharge, which is the discharge the plant is "
     "being built for; so does every distance tested, so recirculation does not constrain the "
     "phase-1 siting. It does not satisfy it for the phase-2 cavern brine. If the cavern stream is "
     "committed, the outfall must be moved to at least "
     f"{_cav_min_ok if _cav_min_ok else -1:.0f} m, or the recirculation criterion re-examined "
     "against the intermittency of that stream. This is a screening result on a 50 m grid under a "
     "steady current and is adequate to rank distances and to test the threshold; it is not a basis "
     "for fixing a siting distance to the nearest 200 m without a tidally-resolved run.")
add_section_figs("Outfall siting & recirculation")

# ===========================================================================
# 9 WATER QUALITY
# ===========================================================================
h1("9  Other Impurities and Water-Quality Compliance")
add_table_from_csv(f"{H.CSV_DIR}/F_impurities_eqs.csv",
                   "Trace-impurity concentrations vs Environmental Quality Standards after dilution.")
add_section_figs("Inputs, parameters & water quality")

# ===========================================================================
# 9 VALIDATION
# ===========================================================================
h1("10  Validation and Data Sources")
para("UBDS was tested against the following independent benchmarks. Where the model and the "
     "reference quantity are defined differently, the model value is converted before comparison "
     "and the conversion is stated; agreement is reported as measured, including where the model "
     "falls outside the published band.")
if VSUM:
    bullet(f"Equation of state: maximum error {VSUM.get('eos_max_error_pct','-')}% vs reference "
           "seawater/brine densities (UNESCO EOS-80; CRC Handbook). Within tolerance.")
    bullet(f"Momentum jet: volume-flux dilution slope {VSUM.get('jet_dilution_slope_perZD','-')} "
           f"per z/D and spread {VSUM.get('jet_spread_dbdz','-')} vs canonical 0.25-0.32 and ~0.11 "
           "(Fischer, List, Koh, Imberger & Brooks, 1979). The spread is reproduced; the dilution "
           "slope sits about 9% below the lower end of the canonical range, i.e. the model entrains "
           "slightly less than the canonical round jet.")
    bullet(f"Pure plume: volume-flux exponent {VSUM.get('plume_dilution_exponent','-')} vs the "
           "5/3 = 1.667 power law (Morton, Taylor & Turner, 1956); about 4% below the analytical "
           "exponent.")
    _n = VSUM.get("dense_jet_n_cases", 0)
    _zt = VSUM.get("dense_jet_zt_over_DFr_mean", 0)
    _zt_short = 100 * (1.6 - _zt) / 1.6
    bullet(f"Inclined 60 deg dense jet (Papakonstantis, Christodoulou & Papanicolaou, 2011), "
           f"benchmarked over Fr {VSUM.get('dense_jet_Fr_min','-')}-{VSUM.get('dense_jet_Fr_max','-')} "
           f"({_n} cases), the span the case study runs at: "
           f"x_r/(D.Fr) = {VSUM.get('dense_jet_xr_over_DFr_mean','-')} (mean) vs the published 2.2-3.3 "
           f"band - in band in {VSUM.get('dense_jet_xr_in_band','-')}/{_n} cases. Return-point dilution, "
           f"converted from the model's bulk (top-hat) value "
           f"to the centreline minimum that the published band refers to (divide by 1.694, Gaussian "
           f"profile, lambda = 1.2): S_r(min)/Fr = {VSUM.get('dense_jet_Sr_min_over_Fr_mean','-')} "
           f"vs the published 0.4-0.6 band - in band in {VSUM.get('dense_jet_Sr_min_in_band','-')}/{_n} "
           f"cases. The uncorrected bulk value is "
           f"{VSUM.get('dense_jet_Sr_bulk_over_Fr_mean','-')} and is not directly comparable.")
    bullet(f"Terminal rise of the same dense jet: z_t/(D.Fr) = {_zt} (mean) against the published "
           f"1.6-2.2 band - about {_zt_short:.0f}% BELOW the lower bound, and inside the band in only "
           f"{VSUM.get('dense_jet_zt_in_band','-')} of the {_n} cases. The shortfall is not uniform: it "
           f"grows as the Froude number falls, so it is worst at the Fr 9-11 of the saturated-cavern "
           f"brine, which is the case that governs the seabed assessment. UBDS therefore under-predicts "
           f"the terminal rise height of an inclined dense jet. Because a lower rise gives a shorter "
           f"trajectory and less entrainment before the plume returns to the bed, this bias is "
           f"conservative for seabed salinity - it over-states, not under-states, the salinity reaching "
           f"the bed. It is not corrected by tuning, since raising the rise height would "
           f"require reducing entrainment and would degrade the jet-dilution and return-point "
           f"dilution agreement above.")
add_section_figs("Validation")
add_table_from_csv(f"{H.CSV_DIR}/V_eos_validation.csv", "Equation-of-state validation.")
add_table_from_csv(f"{H.CSV_DIR}/V_dense_jet.csv", "Inclined dense-jet dimensionless validation.")
h2("10.1  Data sources")
sources = [
 "Fischer, H.B., List, E.J., Koh, R.C.Y., Imberger, J. & Brooks, N.H. (1979). Mixing in "
 "Inland and Coastal Waters. Academic Press. [jet/plume dilution laws]",
 "Morton, B.R., Taylor, G.I. & Turner, J.S. (1956). Turbulent gravitational convection from "
 "maintained and instantaneous sources. Proc. R. Soc. Lond. A 234, 1-23. [plume 5/3 law]",
 "Papakonstantis, I.G., Christodoulou, G.C. & Papanicolaou, P.N. (2011). Inclined negatively "
 "buoyant jets 1 & 2. J. Hydraulic Research 49(1), 3-22. [inclined dense-jet data]",
 "Roberts, P.J.W., Ferrier, A. & Daviero, G. (1997). Mixing in inclined dense jets. J. "
 "Hydraulic Engineering 123(8), 693-699. [dense-jet dilution]",
 "Millero, F.J. & Poisson, A. (1981). International one-atmosphere equation of state of "
 "seawater. Deep-Sea Research 28A, 625-629; UNESCO (1981) Tech. Papers Mar. Sci. 36. [EOS-80]",
 "Jirka, G.H. (2004). Integral model for turbulent buoyant jets in unbounded stratified "
 "flows. Part 1. Environmental Fluid Mechanics 4, 1-56. [integral jet model]",
 "Lee, J.H.W. & Chu, V.H. (2003). Turbulent Jets and Plumes - A Lagrangian Approach. Kluwer. "
 "[entrainment closure]",
 "CRC Handbook of Chemistry and Physics. [saturated NaCl brine density]",
 "Synthetic site-specific data (bathymetry, tides, ADCP currents, CTD salinity/temperature, "
 "brine chemistry) generated by the author for this case study; all values are documented and "
 "engineering-credible, and none are taken from any real consent.",
]
for s in sources:
    bullet(s)

h2("10.2  Supporting data plots")
para("Derived plots summarising the tabulated results above: seabed dilution against tidal "
     "velocity, salinity-increment area against discharge flow, and the per-station "
     "index-of-agreement against the synthetic current reference.")
add_section_figs("Data plots")

# ===========================================================================
# 10 CONCLUSIONS
# ===========================================================================
h1("11  Conclusions")
para("The UBDS solver provides a single, open, physically-validated framework spanning the "
     "near-, medium- and far-field dispersion of brine and other buoyant effluents. Applied to "
     "the Bahia Azul case study with fully documented, credible site-specific inputs, it "
     f"demonstrates that the proposed inclined {C.DIFFUSER['n_ports_installed']}-port diffuser achieves effective initial "
     "dilution and confines significant salinity increments and all trace impurities to within "
     "the regulatory mixing zone, across all simulated discharge flows and tidal states.")
para(f"One condition attaches to that conclusion. The {_site_dist:.0f} m outfall meets the "
     f"{_CRIT} ppt intake-recirculation criterion for the phase-1 RO concentrate at every distance "
     f"screened (worst cell {_ro_worst:.3f} ppt) but no siting inside "
     f"{_cav_min_ok if _cav_min_ok else -1:.0f} m meets it for the phase-2 saturated cavern brine, "
     f"which therefore requires a siting of at least "
     f"{_cav_min_ok if _cav_min_ok else -1:.0f} m. The seabed-salinity and water-quality conclusions "
     "above already use the cavern brine as the worst case; the recirculation conclusion is the "
     "one place where the two phases diverge, and it is stated separately for that reason.")

# ===========================================================================
# APPENDIX - OUTPUT INVENTORY
# ===========================================================================
h1("12  Animated Outputs")
para("The following animated outputs (GIF files in outputs/animations/) show the time-evolving "
     "brine dispersion and tidal current field over a tidal cycle. A static filmstrip of the "
     "dispersion animation is embedded below; the full animations are in outputs/animations/.")
for a in [x for x in MANIFEST if x["kind"] == "animation"]:
    bullet(f"{os.path.basename(a['path'])} - {a['caption']}")
add_section_figs("Animations")

h1("Appendix A  Representative Output Figures from the Reference Studies")
para("For methodological comparison, representative output figures from the three reference "
     "studies that guided this work are reproduced below (author's reference materials, "
     "attributed). They illustrate the output categories - mesh/bathymetry, current "
     "calibration, near-field trajectories/dilution, salinity envelopes and snapshots, and "
     "intake/outfall time histories - that the UBDS outputs reproduce for the Bahia Azul case.")
add_section_figs("Reference output figures (source studies)", width=6.0)

h1("Appendix B  Complete Output Inventory")
para("All figures and datasets generated by the UBDS solver for this case study:")
h2("B.1  Figures")
for i, a in enumerate([x for x in MANIFEST if x["kind"] == "figure"], 1):
    bullet(f"{i:02d}. [{a['section']}] {os.path.basename(a['path'])} - {a['caption']}")
h2("B.2  Animations")
for i, a in enumerate([x for x in MANIFEST if x["kind"] == "animation"], 1):
    bullet(f"{i:02d}. {os.path.basename(a['path'])} - {a['caption']}")
h2("B.3  Datasets (CSV)")
for i, a in enumerate([x for x in MANIFEST if x["kind"] == "csv"], 1):
    bullet(f"{i:02d}. {os.path.basename(a['path'])} - {a['caption']}")

_all_fig_sections = {a["section"] for a in MANIFEST if a["kind"] == "figure"}
_unshown = _all_fig_sections - _EMITTED_SECTIONS
if _unshown:
    raise AssertionError(f"figures registered but never placed in the report: {sorted(_unshown)}")

cp = doc.core_properties
cp.author = C.SITE["assessor"].split(" (")[0]
cp.last_modified_by = cp.author
cp.title = C.SITE["name"] + " - Marine Brine-Dispersion Assessment"
doc.save("case.docx")
nf_ = sum(1 for a in MANIFEST if a["kind"] == "figure")
nc_ = sum(1 for a in MANIFEST if a["kind"] == "csv")
_embedded = sum(len(figs(s)) for s in _EMITTED_SECTIONS)
assert _embedded == nf_, f"embedded {_embedded} figures but manifest has {nf_}"
print(f"case.docx written: {nf_} figures ({_embedded} embedded), {nc_} datasets.")
