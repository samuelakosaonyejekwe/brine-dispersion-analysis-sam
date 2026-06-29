"""
Build answer.docx : how UBDS beats existing models, its novelty, and a
patentability assessment.

Author: Akosa Samuel Onyejekwe (Independent Researcher)
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

INK = RGBColor(0x1B, 0x2A, 0x4A)
ACCENT = RGBColor(0x2A, 0x6F, 0x97)

d = Document()
s = d.styles["Normal"]; s.font.name = "Calibri"; s.font.size = Pt(10.5); s.font.color.rgb = INK
for h in ["Title", "Heading 1", "Heading 2"]:
    try:
        d.styles[h].font.color.rgb = ACCENT if h != "Title" else INK
    except Exception:
        pass


def h1(t): d.add_heading(t, 1)
def h2(t): d.add_heading(t, 2)
def p(t): return d.add_paragraph(t)
def b(t): d.add_paragraph(t, style="List Bullet")
def n(t): d.add_paragraph(t, style="List Number")


t = d.add_heading("", 0)
t.add_run("UBDS - Competitive Advantage, Novelty & Patentability").font.color.rgb = INK
sub = d.add_paragraph()
rr = sub.add_run("Universal Brine Dispersion Solver (UBDS)")
rr.bold = True; rr.font.size = Pt(12); rr.font.color.rgb = ACCENT
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
m = d.add_paragraph(); m.alignment = WD_ALIGN_PARAGRAPH.CENTER
m.add_run("Author: Akosa Samuel Onyejekwe (Independent Researcher)").font.color.rgb = INK
d.add_paragraph("")

# ===========================================================================
h1("1. How UBDS beats other similar existing models")
p("Conventional brine-dispersion assessments require two or three separate, "
  "mostly proprietary tools - a dedicated near-field jet/plume code plus a "
  "far-field hydrodynamic/transport suite - with manual transfer of results "
  "between them. UBDS improves on this in the following concrete ways:")
b("Unified near-to-far field in a single code: removes the manual hand-off and "
  "the mass-conservation error that arises when transferring a plume between two "
  "separate models; the conserved excess-salt flux Q0(S0 - Sa) is carried across "
  "the near-to-far boundary automatically.")
b("Sign-agnostic buoyancy: the same governing equations describe dense brine "
  "(g' > 0), buoyant thermal/freshwater discharges (g' < 0) and neutral tracers. "
  "Most near-field codes are specialised to one buoyancy regime.")
b("Dense-water gravity-current closure in the transport engine: reproduces the "
  "strong bottom-trapping and rapid upward decay of dense brine that 2-D "
  "depth-averaged models structurally cannot represent (they average over depth), "
  "and that full 3-D primitive-equation models obtain only at large computational "
  "cost. This is the headline advantage.")
b("Universal salinity range: the blended equation of state stays valid from fresh "
  "water to saturated brine (~1208 kg/m3 at ~260 psu), whereas standard EOS-80 is "
  "not valid above ~42 psu.")
b("Built-in regulatory-metric engine: mixing-zone radius, salinity-increment "
  "areas, diffusion distances and EQS compliance are computed automatically from "
  "the solution - no separate post-processing tool.")
b("Open, reproducible and validated: physically-standard, documented closures "
  "validated against canonical jet/plume laws and peer-reviewed dense-jet data "
  "(jet spreading 0.114 vs ~0.11; plume volume-flux exponent 1.60 vs 5/3; inclined "
  "dense-jet rise/dilution within the Papakonstantis et al. (2011) bands; "
  "hydrodynamic calibration Willmott d ~ 0.99). Proprietary suites are closed boxes.")
b("Efficient nesting: a coarse hydrodynamic field is interpolated to a fine "
  "transport grid, so a full multi-scenario, multi-tidal-state assessment runs on "
  "an ordinary laptop.")
b("Self-contained design-optimisation and siting workflow: port-angle "
  "optimisation, multi-scenario comparison, and intake-recirculation siting are "
  "all built in, reproducing in one tool what the reference studies did across "
  "several.")

# ===========================================================================
h1("2. All the ways UBDS is novel")
b("Single-framework near-to-far field coupling: an integral near-field jet/plume "
  "model and a hydrodynamic + sigma-layer transport model solved within one "
  "mass-conserving solver, rather than two disconnected codes.")
b("Reduced-gravity-scaled inter-layer gravity-current closure: a buoyancy-driven "
  "vertical settling flux w_s,k = c_s sqrt(max(g'_k, 0) * h_layer) added to the "
  "sigma-layer exchange, giving 3-D-like bottom-trapping at quasi-3D cost. "
  "(Eq. 4.2, model.equations.docx.)")
b("Sign-agnostic formulation: a single buoyancy term (rho_a - rho_e) g handles "
  "dense, buoyant and neutral discharges without code branches - one universal "
  "kernel.")
b("Hypersaline-extended equation of state: a C1 cosine-blended EOS-80 / "
  "brine-anchored density valid across the full 0-300 psu range.")
b("Mass-consistent near-to-far hand-off: the near-field seabed-impact salinity and "
  "footprint set a Dirichlet floor, while the conserved excess-salt mass flux is "
  "injected as a load-scaled bottom-layer source.")
b("Froude-blended entrainment: a local densimetric-Froude weighting that blends "
  "jet and plume entrainment continuously along the trajectory.")
b("Integrated regulatory-compliance engine: mixing-zone, increment-area, "
  "diffusion-distance and EQS-margin metrics derived directly from the solution.")
b("Universal applicability: desalination RO concentrate and saturated "
  "solution-mining cavern brine handled by the same solver and inputs.")

# ===========================================================================
h1("3. Is UBDS patentable? - honest assessment")
p("Note: this is an engineering view of the standard patent criteria, not legal "
  "advice. A qualified patent attorney and a formal prior-art / freedom-to-operate "
  "search are required before filing.")
p("Short answer: possibly - the novel combination and, in particular, the "
  "dense-water gravity-current closure are the realistic candidates, but important "
  "caveats apply.")
h2("3.1 The case for patentability")
b("Novelty / non-obviousness: the specific combination of (a) a sign-agnostic "
  "integral near-field model, (b) a sigma-layer transport model with a "
  "reduced-gravity-scaled inter-layer settling closure, and (c) a mass-conservative "
  "automatic near-to-far coupling, as one solver, does not appear as a single prior "
  "product. The gravity-current closure is the most defensible inventive step.")
b("Utility: a clear industrial application - desalination / salt-cavern brine "
  "outfall design, consenting and environmental impact assessment.")
b("It solves a recognised technical problem: the near-field / far-field two-tool "
  "gap and the inability of 2-D models to capture dense bottom-trapping.")
h2("3.2 The case against (be realistic)")
b("Heavy prior art: integral plume models (Visual Plumes/UM3, CORMIX, "
  "VISJET/JETLAG, CorJet), shallow-water solvers and advection-dispersion transport "
  "are decades-old and well-published. The individual equations (EOS-80, "
  "Morton-Taylor-Turner entrainment, shallow-water, advection-diffusion) are "
  "public-domain and not patentable.")
b("Software / mathematical-method eligibility: in many jurisdictions (e.g. the US "
  "after Alice, and EPC Art. 52 in Europe) an abstract mathematical method or a "
  "computer program 'as such' is not patentable on its own. Eligibility usually "
  "requires tying the method to a technical effect or a specific technical "
  "application/apparatus (e.g. a method of designing or operating a diffuser, or "
  "controlling a discharge, using the solver).")
b("Coefficient calibration alone is not inventive.")
h2("3.3 Most viable routes")
n("Patent the specific technical method - e.g. 'a computer-implemented method for "
  "designing/operating a brine outfall that couples a sign-agnostic integral "
  "near-field model to a sigma-layer transport model via a reduced-gravity "
  "inter-layer gravity-current closure, producing a diffuser-design or "
  "discharge-control output' - framed around the technical effect (improved "
  "dilution prediction leading to diffuser design / operation). The "
  "gravity-current closure plus unified coupling is the inventive core.")
n("Alternatives often stronger for software: copyright (automatic, protects the "
  "code), trade secret (keep the closures proprietary), or defensive publication. "
  "Many numerical solvers are protected this way rather than by patent.")
h2("3.4 Recommendation")
p("Before filing, commission (a) a professional prior-art / freedom-to-operate "
  "search focused on the gravity-current closure and the unified near-to-far "
  "coupling, and (b) a patent attorney to assess subject-matter eligibility in the "
  "target jurisdiction and to frame the claims around the technical application "
  "rather than the mathematics. Realistically, the closure-plus-unified-coupling "
  "method is patentable-leaning; the generic solver as a whole is not.")

d.save("answer.docx")
print("answer.docx written")
